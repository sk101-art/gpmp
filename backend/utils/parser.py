"""
Resume Parser Module

Extracts text from PDF and DOCX files and prepares for LLM extraction.
"""

import fitz
from docx import Document
import io
import logging
import re
from typing import Dict, Any, List

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF."""
    try:
        text = ""
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page_num, page in enumerate(doc):
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page.get_text()
        return text.strip()
    except Exception as e:
        logger.error(f"Error parsing PDF: {str(e)}")
        raise


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        text = ""
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"Error parsing DOCX: {str(e)}")
        raise


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from uploaded file (PDF or DOCX)."""
    if filename.lower().endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename.lower().endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError("Unsupported file format. Please upload PDF or DOCX.")


def preprocess_resume_text(raw_text: str) -> str:
    """Clean and normalize resume text."""
    # Replace multiple newlines with single newline
    text = re.sub(r"\n\n+", "\n", raw_text)
    # Remove extra spaces
    text = re.sub(r" +", " ", text)
    # Normalize whitespace
    text = text.strip()
    return text


def extract_emails(text: str) -> List[str]:
    """Extract email addresses from resume."""
    email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
    return re.findall(email_pattern, text)


def extract_phone_numbers(text: str) -> List[str]:
    """Extract phone numbers from resume."""
    phone_patterns = [
        r"\+?\d{1,3}[-.\s]?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{4,6}", # International/Flexible format
        r"\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}", # US format
    ]
    phones = []
    for pattern in phone_patterns:
        phones.extend(re.findall(pattern, text))
    return list(set(phones))


def extract_urls(text: str) -> List[str]:
    """Extract URLs (LinkedIn, GitHub, portfolio) from resume."""
    url_pattern = r"https?://[^\s]+"
    return re.findall(url_pattern, text)


def prepare_for_extraction(raw_text: str) -> Dict[str, Any]:
    """Prepare resume text for LLM extraction."""
    cleaned_text = preprocess_resume_text(raw_text)
    
    metadata = {
        "emails": extract_emails(raw_text),
        "phone_numbers": extract_phone_numbers(raw_text),
        "urls": extract_urls(raw_text),
        "text_length": len(cleaned_text),
        "line_count": len(cleaned_text.split("\n")),
    }
    
    return {
        "cleaned_text": cleaned_text,
        "metadata": metadata,
    }


def estimate_resume_quality(raw_text: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
    """Estimate resume parsing quality."""
    issues = []
    score = 100
    
    if metadata["text_length"] < 500:
        issues.append("Resume text very short")
        score -= 20
    
    if not metadata["emails"] and not metadata["phone_numbers"]:
        issues.append("No contact information found")
        score -= 10
    
    education_keywords = ["bachelor", "master", "phd", "university", "college", "degree"]
    work_keywords = ["experience", "employment", "work", "responsibilities"]
    
    text_lower = raw_text.lower()
    has_education = any(kw in text_lower for kw in education_keywords)
    has_work = any(kw in text_lower for kw in work_keywords)
    
    if not has_education:
        issues.append("No education section detected")
        score -= 15
    
    if not has_work:
        issues.append("No work experience section detected")
        score -= 15
    
    return {
        "quality_score": max(0, score),
        "issues": issues,
        "is_valid": score > 40,
    }
